from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("outputs/reports/text2sql_current_report.docx")

FONT_CN = "Microsoft YaHei"
FONT_EN = "Calibri"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 35, 45)
MUTED = RGBColor(90, 96, 105)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "EAF4EA"
PALE_GOLD = "FFF4CE"
PALE_RED = "FCE8E6"


def set_run_font(run, size=None, bold=None, color=None, font=FONT_CN):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_EN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_paragraph(doc, text="", style=None, size=11, bold=False, color=INK, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_metric_strip(doc, metrics):
    table = doc.add_table(rows=2, cols=len(metrics))
    table.style = "Table Grid"
    set_table_width(table, [6.5 / len(metrics)] * len(metrics))
    for i, (label, value, fill) in enumerate(metrics):
        set_cell_shading(table.cell(0, i), fill)
        set_cell_shading(table.cell(1, i), fill)
        p0 = table.cell(0, i).paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(label)
        set_run_font(r0, size=8.5, bold=True, color=MUTED)
        p1 = table.cell(1, i).paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(value)
        set_run_font(r1, size=14, bold=True, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, small=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, size=9 if small else 9.5, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 and len(str(value)) <= 16 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=8.5 if small else 9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_note_box(doc, title, body, fill=PALE_GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=INK)
    p.add_run("\n")
    r2 = p.add_run(body)
    set_run_font(r2, size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(11)
        style.font.color.rgb = INK
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.1

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("Execution-Guided Text-to-SQL 后训练阶段报告")
    set_run_font(r, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Text-to-SQL Post-Training Pipeline")
    set_run_font(r, size=9, color=MUTED)


def build_report():
    doc = Document()
    configure_document(doc)

    # Masthead
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("阶段报告")
    set_run_font(r, size=10.5, bold=True, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("基于执行反馈的 Text-to-SQL 后训练与自修复流水线")
    set_run_font(r, size=22, bold=True, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("面向大模型后训练岗位的项目闭环：schema retrieval、SQLite 执行沙盒、self-repair rollout、错误 SQL 采样、DPO preference mining 与 LoRA adapter 评估")
    set_run_font(r, size=11.5, color=MUTED)

    meta = [
        ("项目阶段", "可复现实验闭环已完成；已完成 clean-split DPO 与修复对照"),
        ("基座模型", "Qwen2.5-Coder-3B-Instruct"),
        ("数据集", "Spider train_spider 挖掘偏好数据；Spider dev full 1034 条最终评测"),
        ("报告日期", str(date.today())),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.style = "Table Grid"
    set_table_width(t, [1.3, 5.2])
    for i, (k, v) in enumerate(meta):
        set_cell_shading(t.cell(i, 0), LIGHT_GRAY)
        r0 = t.cell(i, 0).paragraphs[0].add_run(k)
        set_run_font(r0, size=9.5, bold=True, color=INK)
        r1 = t.cell(i, 1).paragraphs[0].add_run(v)
        set_run_font(r1, size=9.5, color=INK)

    doc.add_paragraph()
    add_metric_strip(
        doc,
        [
            ("Base + CM Repair", "68.90%", LIGHT_GRAY),
            ("FilteredDPO + CM Repair", "71.12%", PALE_GREEN),
            ("DPO first-turn", "63.08%", PALE_GREEN),
            ("Executable", "82.32%", PALE_GREEN),
        ],
    )

    add_heading(doc, "1. 项目目标与任务定义", 1)
    add_paragraph(
        doc,
        "本项目将 Text-to-SQL 从一次性生成任务改造成环境交互任务：模型先根据 question + schema 生成 SQL，随后在 SQLite 沙盒中执行，读取语法错误、schema 错误或结果不一致反馈，再进行第 2/3 轮自修复。执行轨迹被进一步用于构造 SFT 样本、DPO 偏好对和 GRPO-style rollout 数据。",
    )
    add_bullet(doc, "核心目标：优化模型生成可执行且执行结果正确的 SQL，而不是只拟合 gold SQL 的 token 分布。")
    add_bullet(doc, "后训练信号：语法合法性、执行正确性、schema 对齐率、结果一致性、repair success 与 attempts penalty。")
    add_bullet(doc, "工程闭环：数据读取 -> schema retrieval -> SQL 生成 -> SQLite 执行 -> feedback -> self-repair -> preference mining -> LoRA DPO -> adapter 评估。")

    doc.add_page_break()
    add_heading(doc, "2. 当前实现范围", 1)
    add_table(
        doc,
        ["模块", "当前状态", "说明"],
        [
            ("Spider/WikiSQL-style 数据读取", "已实现", "支持 JSON/JSONL、db_id + db_root、显式 db_path。"),
            ("SQLite 执行沙盒", "已实现", "只允许 SELECT/WITH，只读连接，带 timeout。"),
            ("Schema retrieval", "已实现", "当前为 token overlap 轻量检索，支持 full/retrieved 对照。"),
            ("Self-repair rollout", "已实现", "max_turns=3，支持 error_only/result_status/oracle_rows 反馈。"),
            ("DPO preference mining", "已实现", "包含 gold_vs_failed、first-turn 对照、true self-repair 三类 pair。"),
            ("错误 SQL 采样器", "已实现", "覆盖 wrong_table、wrong_column、missing_join、wrong_aggregation 等。"),
            ("LoRA DPO / adapter 评估", "已实现", "支持 base model + PEFT adapter 推理评估。"),
            ("SFT / SFT+DPO", "入口已实现", "下一步补完整对照实验。"),
        ],
        [1.55, 1.15, 3.8],
        small=True,
    )

    add_heading(doc, "3. Reward 与偏好数据构造", 1)
    add_paragraph(doc, "当前 rule reward 保留原始设计，并在每个候选 SQL 上记录 rule_reward、cost_reward、schema/executable flags。")
    add_table(
        doc,
        ["条件", "Reward", "含义"],
        [
            ("1 turn 执行结果正确", "1.0", "模型一次生成正确 SQL。"),
            ("第 2/3 轮自修复成功", "0.7 - 0.1*(turn-2)", "鼓励修复成功，同时惩罚更多尝试。"),
            ("SQL 可执行但结果错误", "0.2", "语法/schema 基本可用，但结果不一致。"),
            ("语法或 schema 错误", "-0.2", "不可执行或引用不存在表/列。"),
            ("最大轮数后仍失败", "-0.5", "最终任务失败。"),
        ],
        [2.2, 1.25, 3.05],
        header_fill=LIGHT_BLUE,
    )
    add_paragraph(doc, "DPO pair 当前分为三类：")
    add_bullet(doc, "self_repair_success_vs_failed_attempt：chosen 为修复后正确 SQL，rejected 为第一轮错误 SQL。")
    add_bullet(doc, "gold_vs_failed_attempt：chosen 为 gold SQL，rejected 为第一轮错误 SQL。")
    add_bullet(doc, "first_turn_correct_sample_vs_failed_sample：同一题第一轮采样中，正确 SQL 对比错误 SQL。")

    add_heading(doc, "4. 关键实验结果", 1)
    add_paragraph(doc, "主评估集为 Spider dev full 1034 条，解码为 greedy。Column-Minimal Repair（CM Repair）表示 column-aware execution feedback 加 minimal-SQL repair constraint。")
    add_table(
        doc,
        [
            "Run",
            "First-turn Acc",
            "Final Acc",
            "Repair Success",
            "Schema Align",
            "Executable",
        ],
        [
            ("Base-OneShot", "57.56%", "57.56%", "/", "84.01%", "83.43%"),
            ("Base + CM Repair", "57.56%", "68.90%", "26.71%", "81.64%", "81.20%"),
            ("FilteredDPO-OneShot", "63.08%", "63.08%", "/", "86.24%", "85.95%"),
            ("FilteredDPO + CM Repair", "63.08%", "71.12%", "21.78%", "82.61%", "82.32%"),
        ],
        [1.65, 0.95, 0.9, 1.05, 1.0, 0.95],
        header_fill=LIGHT_BLUE,
        small=True,
    )
    add_note_box(
        doc,
        "核心发现",
        "Filtered DPO 将 first-turn execution accuracy 从 57.56% 提升至 63.08%。在 Column-Minimal Repair 加持下，final execution accuracy 从 Base 的 68.90% 进一步提升到 71.12%。",
        fill=PALE_GREEN,
    )
    add_note_box(
        doc,
        "重要风险与处理",
        "早期 HeavyDPO 使用过量 synthetic preference 后，repair success 大幅提升但 one-shot 稳定性下降。最终采用 reward-margin filtering、SQL edit-distance filtering 与 synthetic 占比控制，得到更稳的 FilteredDPO adapter。",
        fill=PALE_GOLD,
    )

    add_heading(doc, "5. Preference 数据规模与配比", 1)
    add_table(
        doc,
        ["数据来源 / Pair 类型", "数量", "说明"],
        [
            ("Spider 500 自然 rollout pairs", "308", "来自模型真实采样与自修复轨迹。"),
            ("gold_vs_failed_attempt", "119", "失败样本使用 gold SQL 对比首轮错误 SQL。"),
            ("first_turn_correct_sample_vs_failed_sample", "173", "同一题多采样中正确/错误候选对。"),
            ("self_repair_success_vs_failed_attempt", "16", "真正经过执行反馈修复成功的高价值 pair。"),
            ("Synthetic error pairs", "150", "从 1829 条 synthetic 中抽样，避免压过 natural pairs。"),
            ("BalancedDPO 总 pairs", "458", "自然 308 + synthetic 150。"),
        ],
        [2.65, 0.85, 3.0],
        small=True,
    )
    add_paragraph(
        doc,
        "Synthetic error 类型包括 wrong_table、wrong_column、missing_join、wrong_aggregation、wrong_condition、wrong_group_by 和 wrong_order_limit。当前 balanced 配比实验证明：控制 synthetic 占比比单纯扩大偏好数据更重要。",
    )

    add_heading(doc, "6. 错误分析", 1)
    add_paragraph(doc, "开发阶段的 200 条评估轨迹显示，主要失败并不是 schema retrieval 漏表，而是 alias/table-column ownership、join path 和条件 grounding。")
    add_table(
        doc,
        ["错误类型", "候选级数量", "首轮数量", "观察"],
        [
            ("no_such_column", "146", "48", "schema linking 和 join path 仍是主要失败来源。"),
            ("wrong_result", "92", "34", "SQL 可执行但语义或过滤条件不一致。"),
            ("empty_result_or_wrong_result", "22", "9", "常见于条件值、连接路径或 NOT/EXCEPT 类问题。"),
            ("no_such_table", "12", "5", "模型会臆造 song 等不存在表。"),
            ("other_execution_error", "2", "2", "少量执行异常。"),
            ("execution_correct", "112", "102", "包含首轮正确与修复成功候选。"),
        ],
        [1.7, 0.9, 0.8, 3.1],
        small=True,
    )
    add_paragraph(doc, "典型失败案例集中在 concert_singer 数据库：")
    add_bullet(doc, "臆造不存在表：模型生成 `song` 表，但 Spider schema 中实际 song 字段位于 `singer` 表。")
    add_bullet(doc, "列归属错误：例如从 `concert` 别名读取 `Name` / `Capacity`，触发 no such column。")
    add_bullet(doc, "可执行但结果错误：多余 join、错误过滤条件、错误 EXCEPT/NOT IN 语义导致结果集不一致。")

    add_heading(doc, "7. 当前结论与岗位相关性", 1)
    add_paragraph(
        doc,
        "该项目已经形成了一个完整的 execution-feedback 后训练闭环，不只是调用 DPOTrainer，而是覆盖了数据构造、环境反馈、reward 设计、偏好挖掘、LoRA 训练、adapter 评估和错误分析。它与大模型后训练、工具使用 Agent、可验证任务 RLHF/RLAIF 以及多模态 GUI/VQA 中的环境反馈范式高度相似。",
    )
    add_bullet(doc, "后训练能力：构造 chosen/rejected 偏好数据，分析 synthetic preference 过量导致的分布偏移，并通过 balanced DPO 修复。")
    add_bullet(doc, "Agentic workflow：将 SQL 生成转为可执行环境交互任务，模型通过反馈迭代修复。")
    add_bullet(doc, "评估能力：区分 example-level 和 candidate-level 指标，避免采样噪声误读。")
    add_bullet(doc, "工程能力：实现数据流水线、执行沙盒、adapter 加载评估、Markdown 报告和错误诊断脚本。")

    add_heading(doc, "8. 下一步计划", 1)
    add_table(
        doc,
        ["优先级", "任务", "目标产出"],
        [
            ("P0", "补 SFT 与 SFT+DPO 对照", "验证 DPO 是否比纯 SFT 更稳，并形成 Base / SFT / DPO / SFT+DPO 表。"),
            ("P0", "沉淀中文 README 与简历表述", "固定实验结果、命令、错误分析和项目亮点。"),
            ("P1", "增强 schema retrieval", "从 token overlap 升级为 BM25/embedding/schema linking，提高 schema alignment。"),
            ("P1", "细化 reward decomposition", "输出 exec/schema/result/attempt 分项统计。"),
            ("P2", "扩大 Spider 评估规模", "验证 500 条评估集上的泛化趋势。"),
        ],
        [0.8, 2.1, 3.6],
        header_fill=LIGHT_BLUE,
    )

    doc.add_page_break()
    add_heading(doc, "9. 可用于简历的中文表述", 1)
    add_note_box(
        doc,
        "简历 bullet",
        "构建基于执行反馈的 Text-to-SQL 后训练流水线，将 one-shot SQL 生成改造成“生成 SQL-执行-读取错误/结果-自修复”的环境交互任务；结合 schema retrieval、SQLite 沙盒执行、错误 SQL 采样和 execution-guided preference mining 自动构造 DPO 偏好对，并设计语法合法性、执行正确性、schema 对齐率、结果一致性和尝试次数惩罚等 reward 指标。基于 Qwen2.5-Coder-3B 在 Spider train_spider 上挖掘自然偏好对并构造 synthetic hard negatives，通过 reward-margin filtering、SQL edit-distance filtering 和 synthetic 占比控制提升偏好数据质量；在 Spider dev full 1034 clean-split 评测中，Filtered DPO 将 first-turn execution accuracy 从 57.56% 提升至 63.08%，结合 Column-Minimal Repair 后将 final execution accuracy 从 68.90% 提升至 71.12%。",
        fill=LIGHT_GRAY,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()
